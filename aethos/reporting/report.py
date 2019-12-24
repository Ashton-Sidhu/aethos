import os
import platform
import subprocess
from datetime import datetime
from aethos.config import cfg, DEFAULT_REPORT_DIR, DEFAULT_IMAGE_DIR
from aethos.config.config import _global_config
from aethos.util import _make_dir
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
import pandas as pd


class Report:
    def __init__(self, report_name):

        if not cfg["report"]["dir"]:  # pragma: no cover
            report_dir = DEFAULT_REPORT_DIR
        else:
            report_dir = cfg["report"]["dir"]

        _make_dir(report_dir)

        self.report_name = report_name
        self.docx = _global_config["word_report"]

        if os.path.exists(self.report_name):
            self.filename = self.report_name
            self.docx_filename = self.filename.replace(".txt", ".docx")

            if self.docx:  # pragma: no cover
                self.doc = Document(docx=self.docx_filename)
                style = self.doc.styles
                if not style["Indent"]:
                    style.add_style("Indent", WD_STYLE_TYPE.PARAGRAPH)
                    paragraph_format = style.paragraph_format
                    paragraph_format.left_indent = Inches(0.25)
        else:
            self.filename = "{}/{}-{}".format(
                report_dir, report_name, datetime.now().strftime("%d-%m-%Y_%I-%M-%S%p")
            )

            if self.docx:  # pragma: no cover
                self.doc = Document()
                style = self.doc.styles.add_style("Indent", WD_STYLE_TYPE.PARAGRAPH)
                paragraph_format = style.paragraph_format
                paragraph_format.left_indent = Inches(0.25)
                self.docx_filename = self.filename + ".docx"

            self.filename += ".txt"

            self.report_environtment()

        if not cfg["images"]["dir"]:
            self.image_dir = DEFAULT_IMAGE_DIR
        else:
            self.image_dir = cfg["images"]["dir"]

    def write_header(self, header: str, level=1):
        """
        Writes a header to a report file
        
        Parameters
        ----------
        header : str
            Header for a report file
            Examples: 'Cleaning', 'Feature Engineering', 'Modelling', etc.
        """

        write = True

        if os.path.exists(self.filename):
            with open(self.filename, "r") as f:
                if header in f.read():
                    write = False

        if write:
            if self.docx:  # pragma: no cover
                self.doc.add_heading(header, level=level)
                self.doc.save(self.docx_filename)

            with open(self.filename, "a+") as f:
                f.write("\n" + header + "\n\n")

    def write_contents(self, content: str):
        """
        Write report contents to report.
        
        Parameters
        ----------
        content : str
            Report Content
        """

        write = True

        if os.path.exists(self.filename):
            with open(self.filename, "r") as f:
                if content in f.read():
                    write = False

        if write:
            if self.docx:  # pragma: no cover
                self.doc.add_paragraph(content, style="Indent")
                self.doc.save(self.docx_filename)

            with open(self.filename, "a+") as f:
                f.writelines(content)

    def report_technique(self, technique: str, list_of_cols=[]):
        """
        Writes analytic technique info to the report detailing what analysis was run
        and why it was ran.
        
        Parameters
        ----------
        technique : str
            Analytic technique
            
        list_of_cols : list
            List of columns
        """

        if list(list_of_cols) != []:
            for col in list_of_cols:
                log = technique.replace("column_name_placeholder", col)

                self.write_contents(log)

        else:
            self.write_contents(technique)

    def log(self, log: str):
        """
        Logs info to the report file.
        
        Parameters
        ----------
        log : str
            Information to write to the report file
        """

        self.write_contents(log + "\n")

    def report_gridsearch(self, clf, verbose):
        """
        Logs gridsearch results.
        """

        content = f"Tuning hyper-parameters:\n\n \
            Best parameters set found on training set:\n \
                {clf.best_params_}\n\n \
            Grid scores on training set:\n\n"

        self.write_contents("\n".join(map(str.lstrip, content.split("\n"))))

        if verbose:
            print(content)

        means = clf.cv_results_["mean_test_score"]
        stds = clf.cv_results_["std_test_score"]
        for mean, std, params in zip(means, stds, clf.cv_results_["params"]):
            stats = "{:.2%} (+/-{}) for {}\n".format(float(mean), std * 2, params)
            self.write_contents(stats)
            if verbose:
                print(stats)

            self.write_contents("\n")

    def report_classification_report(self, report):
        """
        Logs classification report results.
        """

        self.write_contents("Detailed classification report:\n")
        self.write_contents("The model is trained on the full training set.")
        self.write_contents("The scores are computed on the full test set.\n\n")
        self.write_contents(report)

    def report_environtment(self):
        """
        Reports local environment info.
        """

        system = platform.system()

        self.write_header("Environment")

        self.write_contents("OS Version          : {}\n".format(platform.platform()))
        self.write_contents("System              : {}\n".format(system))
        self.write_contents("Machine             : {}\n".format(platform.machine()))

        if system == "Windows":  # pragma: no cover
            memory = subprocess.check_output(
                "wmic memorychip get capacity", shell=True, universal_newlines=True
            )
            total_mem = 0
            for m in memory.split("  \r\n")[1:-1]:
                total_mem += int(m)

            self.write_contents(
                "Processor           : {}\n".format(platform.processor())
            )
            self.write_contents(
                "Memory              : {:.2f} GB\n".format(
                    round(total_mem / (1024 ** 2))
                )
            )
        elif platform.system() == "Darwin":  # pragma: no cover
            processor = subprocess.check_output(
                "/usr/sbin/sysctl -n machdep.cpu.brand_string",
                shell=True,
                universal_newlines=True,
            )
            memory = subprocess.check_output(
                "sysctl -n hw.memsize", shell=True, universal_newlines=True
            )
            self.write_contents("Processor           : {}\n".format(processor.strip()))
            self.write_contents(
                "Memory              : {}\n".format(
                    round(int(memory.strip()) / (1024 ** 2))
                )
            )
        elif platform.system() == "Linux":
            processor = subprocess.check_output(
                "cat /proc/cpuinfo | grep 'model name'",
                shell=True,
                universal_newlines=True,
            )
            with open("/proc/meminfo", "r") as f:
                lines = f.readlines()

            memory = [int(s) for s in lines[0].split() if s.isdigit()][0]
            processor = processor.replace("\t", "").split("model name:")[1].strip()
            self.write_contents("Processor           : {}\n".format(processor))
            self.write_contents(
                "Memory              : {:.2f} GB\n".format(round(memory / (1024 ** 2)))
            )

        self.write_contents("CPU Count           : {}\n".format(os.cpu_count()))
        self.write_contents(
            "Python Version      : {}\n".format(platform.python_version())
        )
        self.write_contents(
            "Python Compiler     : {}\n".format(platform.python_compiler())
        )
        self.write_contents(
            "Python Build        : {}\n".format(platform.python_build())
        )

    def write_image(self, name):  # pragma: no cover

        if self.docx:
            self.doc.add_picture(
                os.path.join(self.image_dir, name), width=Inches(6), height=Inches(4)
            )

            self.doc.save(self.docx_filename)

    def write_metrics(self, df: pd.DataFrame):

        self.write_header("Metrics", level=2)
        self.write_contents(df.to_string())

        if self.docx:  # pragma: no cover
            t = self.doc.add_table(df.shape[0] + 1, df.shape[1])

            t.style = "Light List Accent 1"

            # add the header rows.
            for j in range(df.shape[-1]):
                t.cell(0, j).text = df.columns[j]

            # add the rest of the data frame
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i + 1, j).text = str(df.values[i, j])

            self.doc.save(self.docx_filename)
